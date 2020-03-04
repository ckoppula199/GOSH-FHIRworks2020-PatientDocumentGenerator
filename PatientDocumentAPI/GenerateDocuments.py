from DataRetrieval import *
import docx
from docx.shared import Inches
from io import BytesIO
from dateutil.parser import parse
import matplotlib.pyplot as plt


def add_response_line(paragraph, run_number):
    """
    Adds a line underneath a question for patients to write their answer
    :param paragraph: Paragraph to add line to
    :param run_number: run number of paragraph to add page breaks to
    """
    paragraph.runs[run_number].add_break()
    paragraph.runs[run_number].add_break()
    paragraph.add_run("________________________________________________________")


def add_address(feedback_data, doc):
    """
    Adds address of the patient to the top right of the page
    :param feedback_data: dictionary containing data on the patient from which address data is extracted.
    :param doc: Document to add address to
    """
    address_line, city, state, postcode, country = get_patient_address_data(feedback_data)
    address = doc.add_paragraph()
    for line in address_line:
        address.add_run(line).add_break()
    address.add_run(city).add_break()
    address.add_run(state).add_break()
    address.add_run(postcode).add_break()
    address.add_run(country).add_break()
    address.alignment = 2


def add_greeting_with_name(feedback_data, doc):
    """
    Adds greeting with the name of the patient to the page
    :param feedback_data: dictionary containing data on the patient from which name data is extracted.
    :param doc: Document to add address to
    """
    prefix, first_name, last_name = get_patient_name_data(feedback_data)
    doc.add_paragraph('Hello ' + prefix + " " + first_name + " " + last_name + ",")


class FeedbackForm():
    """
    Class generates a word document asking the Patient for the hospital and clinic they visited as well as how likely
    they are to recommend the service they received to someone else and for a comment explaining their answer. This can
    then be emailed or printed out and sent to the patient to collect patient service feedback data.
    """

    def __init__(self, patient_id, feedback_data):
        self.doc = docx.Document()
        self.patient_id = patient_id
        self.feedback_data = feedback_data
        self.intro, self.recommendation, self.comment, self.hospital, self.clinic = get_feedback_text(feedback_data)

    def generate_feedback_form(self):
        """
        Creates the Word document asking the patient for feedback.
        """
        add_address(self.feedback_data, self.doc)
        self.doc.add_paragraph('Patient Feedback Form', 'Title')
        add_greeting_with_name(self.feedback_data, self.doc)
        welcome = self.doc.add_paragraph(self.intro)
        ask_for_hospital = self.doc.add_paragraph(self.hospital)
        add_response_line(ask_for_hospital, 0)
        ask_for_clinic = self.doc.add_paragraph(self.clinic)
        add_response_line(ask_for_clinic, 0)
        ask_for_recommendation = self.doc.add_paragraph(self.recommendation)
        ask_for_recommendation.add_run("Please tick your choice from the options below.")
        ask_for_recommendation.runs[0].add_break()
        self.add_feedback_table()
        self.add_comment_box(self.comment)
        self.doc.save(self.patient_id + " feedback request.docx")

    def add_feedback_table(self):
        """
        Adds table of response options on how likely patient is to recommend the service they received
        """
        table = self.doc.add_table(rows=2, cols=4)
        row1 = table.rows[0]
        row2 = table.rows[1]
        row1.cells[0].text = "Extremely Likely"
        row1.cells[1].text = "Likely"
        row1.cells[2].text = "Unlikely"
        row1.cells[3].text = "Extremely Unlikely"
        for index in range(4):
            row1.cells[index].paragraphs[0].runs[0].font.bold = True
            row1.cells[index].paragraphs[0].alignment = 1
        row2.cells[0].text = "\n"
        table.style = self.doc.styles['Table Grid']

    def add_comment_box(self, comment):
        """
        Adds line asking for a comment on service received and adds a box for the patient to add their comment to
        :param comment: Line of text asking patient to provide a comment
        """
        ask_for_comment = self.doc.add_paragraph("\n" + comment)
        answer_box = self.doc.add_table(rows=1, cols=1)
        answer_box.rows[0].cells[0].text = '\n\n\n\n\n\n\n\n\n\n'
        answer_box.style = self.doc.styles['Table Grid']


class PatientHealthForm():

    def __init__(self, patient_id, name, patient_data):
        self.doc = docx.Document()
        self.patient_id = patient_id
        self.name = name
        self.patient_data = patient_data

    def generate_patient_data_form(self):
        """
        Creates a word document containing visualisations of patient health data regarding Weight, BMI, Heart rate
        and Respiratory rate
        """
        self.doc.add_paragraph('Patient Health Data', 'Title')
        self.doc.add_paragraph("Name: " + self.name)
        self.doc.add_paragraph("ID: " + self.patient_id)
        self.doc.add_paragraph("This document contains graphs detailing changes in the patients vital signs over time\n")

        for data_type in self.patient_data.keys():
            dates = self.patient_data[data_type]['Dates']
            values = self.patient_data[data_type]['Values']
            unit = self.patient_data[data_type]['Unit']
            dates = [parse(date) for date in dates]

            memfile = BytesIO()
            plt.figure()
            plt.plot_date(dates, values, color='orange')
            plt.xlabel('Date')
            plt.ylabel(unit)
            plt.title(data_type)
            plt.savefig(memfile)
            self.doc.add_paragraph(data_type + " over time").runs[0].bold = True
            self.doc.add_picture(memfile, height=Inches(3))

        self.doc.save(self.patient_id + " health data.docx")


class PatientDataForm():

    def __init__(self, patient):
        self.doc = docx.Document()
        self.patient = patient

    def generate_patient_info_form(self):
        """
        Creates a word document with all of the patients personal information and asks them to check the details.
        Provides lines below the details for the patients to update the data if any of it is wrong or out of date.
        """
        para_1 = self.doc.add_paragraph('Patient Details', 'Title')
        para_2 = self.doc.add_paragraph("This document contains the personal information pertaining to " + self.patient.full_name())
        para_2.add_run("\nIf any details are incorrect or out of date please write the correct value on the line below "
                       "the detail")
        para_3 = self.doc.add_paragraph("ID: " + self.patient.uuid)
        add_response_line(para_3, 0)
        para_4 = self.doc.add_paragraph("First Name: " + self.patient.name.given)
        add_response_line(para_4, 0)
        para_5 = self.doc.add_paragraph("Surname Name: " + self.patient.name.family)
        add_response_line(para_5, 0)
        para_6 = self.doc.add_paragraph("Prefix: " + self.patient.name.prefix)
        add_response_line(para_6, 0)
        para_7 = self.doc.add_paragraph("First Name: " + self.patient.name.given)
        add_response_line(para_7, 0)
        para_8 = self.doc.add_paragraph("Gender: " + self.patient.gender)
        add_response_line(para_8, 0)
        para_9 = self.doc.add_paragraph("Date of birth: " + str(self.patient.birth_date))
        add_response_line(para_9, 0)
        para_10 = self.doc.add_paragraph("Address: " + self.patient.addresses[0].lines[0])
        add_response_line(para_10, 0)
        para_11 = self.doc.add_paragraph("City: " + self.patient.addresses[0].city)
        add_response_line(para_11, 0)
        para_12 = self.doc.add_paragraph("State: " + self.patient.addresses[0].state)
        add_response_line(para_12, 0)
        para_13 = self.doc.add_paragraph("Postal Code: " + self.patient.addresses[0].postal_code)
        add_response_line(para_13, 0)
        para_14 = self.doc.add_paragraph("Country: " + self.patient.addresses[0].country)
        add_response_line(para_14, 0)
        para_15 = self.doc.add_paragraph("Marital Status: " + str(self.patient.marital_status))
        add_response_line(para_15, 0)
        para_16 = self.doc.add_paragraph("Main language: " + self.patient.communications.languages[0])
        add_response_line(para_16, 0)
        para_17 = self.doc.add_paragraph("Driving License Number: " + self.patient.get_identifier('DL'))
        add_response_line(para_17, 0)
        para_18 = self.doc.add_paragraph("Social Security Number: " + self.patient.get_identifier("SS"))
        add_response_line(para_18, 0)
        self.doc.save(self.patient.uuid + " details.docx")




# 8f789d0b-3145-4cf2-8504-13159edaa747
